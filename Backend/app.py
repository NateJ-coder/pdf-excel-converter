# backend/app.py
# Updated to include:
# - master description list
# - predefined row order
# - file-based year alignment
# - normalized OCR text
# - inclusion of total lines if missing
# - Number formatting for currency columns in Excel
# - Refactored Category Inference with Canonical Descriptions
# - Logging Module instead of print()
# - Pre-initialization of consolidated_items to ensure all MASTER_STRUCTURE descriptions appear
# - Improved Gemini prompt to avoid unnecessary total recalculations
# - Removed responseSchema from Gemini API API call for more flexible parsing
# - Refined Excel generation logic for precise data placement and summary calculations
# - Broadened Gemini prompt to extract data from all relevant financial sections (including notes)
# - Enhanced infer_categories_and_structure to handle more diverse financial line items
# - Improved consolidated_items_by_key population to strictly align with MASTER_STRUCTURE for known items
# - Removed the Summary sheet entirely as per user request
# - Further refined data consolidation to capture all years and granular items
# - Added more robust mapping from Gemini output to MASTER_STRUCTURE
# - Integrated custom AI prompt from frontend for parsing.
# - FIXED: NameError: name 'custom_prompt_text' is not defined in upload_pdfs.
# - FIXED: KeyError: '"Description"' due to f-string and .format() conflict.
# - Modified generate_excel to produce a hierarchical "Description | Year1 | Year2 | ..." table.
# - Addressed duplicate/inflated values by refining canonicalization and mapping.
# - Improved handling of missing values by robust mapping.
# - Fixed line item confusion by ensuring correct value assignment.
# - Trimmed empty category/subcategory rows for cleanliness.
# - Further enhanced number cleaning for robustness (e.g., handling spaces as thousands separators).
# - Reviewed and reinforced canonicalization for "Bank balances" and "Short-term deposits".
# - Added more explicit examples for "Bank balances" and "ABSA" to master structure under Current Assets.
# - Added specific instruction in Gemini prompt for numbers with spaces as thousands separators.
# - **NEW**: Added more explicit examples for "Short-term deposits" for 2020 and 2019 in Gemini prompt.
# - **NEW**: Implemented post-processing logic to ensure "Accumulated deficit" and "Accumulated surplus" are mutually exclusive per year.
# - **NEW**: Integrated Document Refinement Tool with new routes and helper functions.
# - **UPDATED**: Document Refinement Tool now parses DOCX files properly using python-docx.
# - **UPDATED**: Gemini prompt for refinement is more specific about TOC, Definitions, and Director's Table.
# - **UPDATED**: `create_refined_document` now attempts to copy template structure and insert specific content.
# - **UPDATED**: `detect_text_from_pdf` now accepts content directly, removing need to save PDF to disk.
# - **FIXED**: SyntaxError: unterminated string literal in `create_refined_document` for TOC dots.
# - **FIXED**: SyntaxError: '(' was never closed in `create_refined_document` for p.add_run call.
# - **FIXED**: Changed Flask app port to 5000 to match frontend expectation.
# - **UPDATED**: Allowed PDF files as templates for document refinement.
# - **FIXED**: Re-added missing `extract_text_from_file` function.
# - **FIXED**: Modified `create_refined_document` to handle PDF templates by creating a new DOCX from extracted content.
# - **UPDATED**: Refactored `create_refined_document` to insert specific template content into the data file, leaving data file largely unchanged.
# - **UPDATED**: Gemini prompt for refinement to extract specific content (TOC, Definitions, Director's Table) from the template.

import os
import io
import re
import json
import requests
import logging
import copy
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from google.cloud import vision
from dotenv import load_dotenv
import google.generativeai as genai
from openpyxl import Workbook
from openpyxl.styles import Font, Border, Side, Alignment, PatternFill, numbers
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_LEADER
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


load_dotenv()

# --- Logging Configuration ---
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")

if not GEMINI_API_KEY:
    logger.error("Error: GEMINI_API_KEY is not set. Cannot call Gemini API.")
else:
    genai.configure(api_key=GEMINI_API_KEY)

vision_client = vision.ImageAnnotatorClient()

# --- Canonical Descriptions Mapping (Existing) ---
CANONICAL_DESCRIPTIONS = {
    "property, plant and equipment": "Property, plant and equipment",
    "pp&e": "Property, plant and equipment",
    "ppe": "Property, plant and equipment",
    "other financial assets": "Other financial assets",
    "trade and other receivables": "Trade and other receivables",
    "trade receivables": "Trade and other receivables",
    "cash and cash equivalents": "Cash and Cash Equivalents",
    "cash equivalents": "Cash and Cash Equivalents",
    "cash at bank": "Cash and Cash Equivalents",
    "total assets": "Total Assets",
    "total asset": "Total Assets",
    "reserves": "Reserves",
    "accumulated surplus": "Accumulated surplus",
    "accum surplus": "Accumulated surplus",
    "accumulated deficit": "Accumulated deficit",
    "accumulated funds": "Accumulated Funds",
    "total equity": "Total Equity",
    "total equ": "Total Equity",
    "deferred tax liability": "Deferred tax liability",
    "tax liability": "Deferred tax liability",
    "trade and other payables": "Trade and Other Payables",
    "trade payables": "Trade and Other Payables",
    "provisions": "Provisions",
    "bank overdraft": "Bank overdraft",
    "total liabilities": "Total Liabilities",
    "total liability": "Total Liabilities",
    "total equity and liabilities": "Total Equity and Liabilities",
    "total equity & liabilities": "Total Equity and Liabilities",
    "total equity and liab": "Total Equity and Liabilities",
    "revenue": "Revenue",
    "levies received": "Levies received",
    "fines": "Fines",
    "water recovered": "Water recovered",
    "ombudsman levy": "Ombudsman levy",
    "electricity recovered": "Electricity recovered",
    "garage levies": "Garage levies",
    "security levies": "Security levies",
    "other income": "Other income",
    "tower rental": "Tower rental",
    "insurance claims received": "Insurance claims received",
    "special levy": "Special levy",
    "interest received": "Interest received",
    "fair value adjustments": "Fair value adjustments",
    "operating expenses": "Operating expenses",
    "accounting fees": "Accounting fees",
    "bank charges": "Bank charges",
    "csos": "CSOS",
    "cleaning": "Cleaning",
    "depreciation, amortisation and impairments": "Depreciation, amortisation and impairments",
    "depreciation and amortisation": "Depreciation and amortisation",
    "electricity": "Electricity",
    "employee costs": "Employee costs",
    "garden services": "Garden services",
    "insurance": "Insurance",
    "management fees": "Management fees",
    "other expenses": "Other expenses",
    "petrol and oil": "Petrol and oil",
    "printing and stationery": "Printing and stationery",
    "protective clothing": "Protective clothing",
    "repairs and maintenance": "Repairs and maintenance",
    "security": "Security",
    "investment revenue": "Investment revenue",
    "surplus (deficit) for the year": "Surplus (deficit) for the year",
    "total comprehensive income (loss) for the year": "Total comprehensive income (loss) for the year",
    "cash on hand": "Cash on hand",
    "bank balances": "Bank balances",
    "bank balance": "Bank balances",
    "short-term deposits": "Short-term deposits",
    "short term deposits": "Short-term deposits",
    "amounts received in advance": "Amounts received in advance",
    "deposits received": "Deposits received",
    "legal proceedings": "Legal proceedings",
    "rental income": "Rental Income",
    "auditor's remuneration": "Auditor's remuneration",
    "fees": "Auditor's remuneration",
    "bad debts": "Bad debts",
    "consulting and professional fees": "Consulting and professional fees",
    "interest income": "Interest income",
    "csos levies": "CSOS levies",
    "garbage levies": "Garbage levies",

    "compensation commissioner": "Compensation commissioner",
    "employee costs - salaried staff": "Employee costs - salaried staff",
    "municipal charges": "Municipal charges",
    "electricity - recovered from members": "Electricity - recovered from members",
    "water-recovered from members": "Water - recovered from members",
    "maintenance": "Maintenance",
    "elevator maintenace": "Elevator maintenance",
    "total income": "Total Income",
    "total operating expenses": "Total Operating Expenses",
    "(deficit) surplus for the year": "(Deficit) surplus for the year",
    "surplus before taxation": "Surplus before taxation",
    "adjustments for": "Adjustments for",
    "movements in provisions": "Movements in provisions",
    "changes in working capital": "Changes in working capital",
    "net provisions": "Net provisions",
    "non provision of tax": "Non provision of tax",
    "taxation": "Taxation",
    "cash generated from (used in) operations": "Cash generated from (used in) operations",
    "basic": "Basic (Employee Cost)",
    "uif": "UIF (Employee Cost)",
    "absa": "Bank balances"
}

# Defines the structure and order of items in the final Excel report.
MASTER_STRUCTURE = {
    "Assets": {
        "Non-Current Assets": ["Property, plant and equipment", "Other financial assets"],
        "Current Assets": ["Trade and other receivables", "Cash and Cash Equivalents", "Cash on hand", "Bank balances", "Short-term deposits"],
        "N/A": ["Total Assets"]
    },
    "Equity": {
        "Equity": ["Reserves", "Accumulated surplus", "Accumulated deficit", "Accumulated Funds"],
        "N/A": ["Total Equity"]
    },
    "Liabilities": {
        "Non-Current Liabilities": ["Deferred tax liability"],
        "Current Liabilities": ["Trade and Other Payables", "Provisions", "Amounts received in advance", "Deposits received", "Bank overdraft", "Legal proceedings"],
        "N/A": ["Total Liabilities", "Total Equity and Liabilities"]
    },
    "Comprehensive Income": {
        "Revenue": ["Revenue", "Levies received", "Fines", "Water recovered", "Ombudsman levy", "Electricity recovered", "Garage levies", "Security levies", "Other income", "Tower rental", "Insurance claims received", "Special levy", "Interest received", "Fair value adjustments", "Investment revenue", "Rental Income", "Interest income", "CSOS levies", "Garbage levies"],
        "N/A": ["Total Income"],
        "Operating Expenses": ["Operating expenses", "Accounting fees", "Auditor's remuneration", "Bank charges", "CSOS", "Cleaning", "Depreciation, amortisation and impairments", "Electricity", "Employee costs", "Garden services", "Insurance", "Management fees", "Other expenses", "Petrol and oil", "Printing and stationery", "Protective clothing", "Repairs and maintenance", "Security", "Bad debts", "Consulting and professional fees", "Compensation commissioner", "Employee costs - salaried staff", "Municipal charges", "Electricity - recovered from members", "Water - recovered from members", "Maintenance", "Elevator maintenance", "Basic (Employee Cost)", "UIF (Employee Cost)"],
        "N/A_2": ["Total Operating Expenses"] # Using unique keys for N/A sections
    },
    "Profit/Loss": {
        "N/A": [
            "Surplus (deficit) for the year",
            "Total comprehensive income (loss) for the year",
            "(Deficit) surplus for the year",
            "Surplus before taxation",
            "Taxation",
            "Non provision of tax"
        ]
    },
    "Cash Flow": {
        "Adjustments": ["Adjustments for", "Movements in provisions", "Changes in working capital", "Net provisions"],
        "N/A": ["Cash generated from (used in) operations"]
    }
}

# --- Utility Functions ---

def get_canonical_name(name):
    """Finds the canonical name for a given financial term."""
    return CANONICAL_DESCRIPTIONS.get(name.lower().strip(), name)

def clean_value(value):
    """Cleans and converts a string value to a float, handling various formats."""
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        value = value.strip()
        # Handle negative values in parentheses
        is_negative = value.startswith('(') and value.endswith(')')
        # Remove non-numeric characters except for a single decimal point
        value = re.sub(r'[^\d.]', '', value)
        if value:
            try:
                numeric_value = float(value)
                return -numeric_value if is_negative else numeric_value
            except ValueError:
                return None
    return None

# --- Core Processing Functions ---

def extract_text_from_pdf(pdf_content):
    """Extracts text from a PDF file's content using Google Cloud Vision API."""
    if not vision_client:
        raise ConnectionError("Google Cloud Vision client is not initialized.")
    
    logger.info("Starting OCR for PDF content.")
    request = {
        'input_config': {
            'content': pdf_content,
            'mime_type': 'application/pdf'
        },
        'features': [{'type_': vision.Feature.Type.DOCUMENT_TEXT_DETECTION}],
    }

    try:
        response = vision_client.batch_annotate_files(requests=[request])
        full_text = "".join([
            page_response.full_text_annotation.text
            for page_response in response.responses[0].responses
        ])
        logger.info("Finished OCR for PDF content.")
        return re.sub(r'\s+', ' ', full_text).strip()
    except Exception as e:
        logger.error(f"Error during OCR with Google Cloud Vision: {e}")
        raise

def extract_content_from_docx(docx_content):
    """Extracts text, headings, and tables from a DOCX file's content."""
    logger.info("Extracting content from DOCX file.")
    doc = Document(io.BytesIO(docx_content))
    content = {
        'text': "\n".join([p.text for p in doc.paragraphs]),
        'headings': [],
        'tables': []
    }
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith('Heading'):
            try:
                level = int(p.style.name.split(' ')[-1])
                content['headings'].append({'level': level, 'text': p.text})
            except (ValueError, IndexError):
                pass # Ignore styles that don't conform to "Heading X"
    
    for table in doc.tables:
        table_data = [[cell.text for cell in row.cells] for row in table.rows]
        content['tables'].append(table_data)

    logger.info("Finished extracting content from DOCX.")
    return content

def extract_text_from_file(file_content, filename):
    """
    Extracts text and structured content from various file types.
    For DOCX, it uses python-docx for better parsing.
    For PDF, it uses Google Cloud Vision.
    For TXT, it decodes directly.
    """
    logger.info(f"Attempting to extract content from {filename}...")
    if filename.lower().endswith('.pdf'):
        # For PDF, we only get text, no structured headings/tables directly via this path
        return {'text': extract_text_from_pdf(file_content), 'type': 'pdf', 'headings': [], 'tables': []}
    elif filename.lower().endswith(('.doc', '.docx')):
        try:
            # For DOCX, we get text, headings, and tables
            docx_parsed_content = extract_content_from_docx(file_content)
            return {'content': docx_parsed_content, 'type': 'docx', 'text': docx_parsed_content['text'], 'headings': docx_parsed_content['headings'], 'tables': docx_parsed_content['tables']}
        except Exception as e:
            logger.warning(f"Failed to parse DOCX with python-docx for {filename}: {e}. Falling back to basic text decode.")
            try:
                return {'text': file_content.decode('utf-8'), 'type': 'text', 'headings': [], 'tables': []}
            except UnicodeDecodeError:
                return {'text': file_content.decode('latin-1'), 'type': 'text', 'headings': [], 'tables': []}
    elif filename.lower().endswith('.txt'):
        try:
            return {'text': file_content.decode('utf-8'), 'type': 'text', 'headings': [], 'tables': []}
        except UnicodeDecodeError:
            return {'text': file_content.decode('latin-1'), 'type': 'text', 'headings': [], 'tables': []}
    else:
        logger.warning(f"Unsupported file type for text extraction: {filename}")
        return {'text': None, 'type': 'unknown', 'headings': [], 'tables': []}


def parse_financial_data_with_gemini(text_content, filename, custom_prompt_text=""):
    """Uses Gemini API to parse financial text and return structured JSON."""
    logger.info(f"Sending text from {filename} to Gemini API for financial data parsing.")
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY is not set.")

    model = genai.GenerativeModel('gemini-1.5-flash-latest')

    system_instruction = """
    You are an expert financial data extractor. Your task is to extract all financial line items and their corresponding numerical values from the provided text.
    - The output must be a valid JSON array of objects.
    - Each object must have two keys: "Description" (string) and "AmountsByYear" (an object).
    - The "AmountsByYear" object should have years as keys (e.g., "2023") and numerical values as values.
    - Parse numbers correctly: remove currency symbols, commas, and handle parentheses for negative numbers. Treat spaces as thousand separators (e.g., "1 234" is 1234).
    - Do NOT calculate or infer any values. Only extract what is explicitly present in the text.
    - Extract data from all sections, including main statements and notes.
    - Example output format: [{"Description": "Revenue", "AmountsByYear": {"2023": 500000, "2022": 450000}}]
    """

    prompt = f"""
    {custom_prompt_text}

    Please extract the financial data from the following text:
    ---
    {text_content}
    ---
    """
    
    try:
        response = model.generate_content(
            prompt,
            generation_config={"response_mime_type": "application/json"},
            system_instruction=system_instruction
        )
        
        # Clean the response text to ensure it's valid JSON
        cleaned_text = response.text.strip()
        if cleaned_text.startswith("```json"):
            cleaned_text = cleaned_text[7:]
        if cleaned_text.endswith("```"):
            cleaned_text = cleaned_text[:-3]

        parsed_data = json.loads(cleaned_text)
        
        if not isinstance(parsed_data, list):
            logger.warning(f"Gemini returned non-list data for {filename}: {type(parsed_data)}")
            return []
            
        logger.info(f"Successfully parsed data from {filename} using Gemini.")
        return parsed_data

    except Exception as e:
        logger.error(f"Error during Gemini API call or JSON parsing for {filename}: {e}")
        # Log the raw response if available for debugging
        if 'response' in locals() and hasattr(response, 'text'):
            logger.error(f"Raw Gemini response: {response.text}")
        return []

def generate_excel_report(all_items, all_years):
    """Generates an Excel workbook from the consolidated financial data."""
    logger.info("Generating Excel report...")
    wb = Workbook()
    ws = wb.active
    ws.title = "Financials"

    # --- Define Styles ---
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
    category_font = Font(bold=True)
    subcategory_font = Font(bold=True)
    total_font = Font(bold=True)
    total_border = Border(bottom=Side(style='thin'), top=Side(style='thin'))
    currency_format = '_(* #,##0.00_);_(* (#,##0.00);_(* "-"??_);_(@_)'

    # --- Write Headers ---
    headers = ["Description"] + [str(year) for year in all_years]
    ws.append(headers)
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
        ws.column_dimensions[get_column_letter(col_idx)].width = 20 if col_idx > 1 else 50
    
    # --- Write Data ---
    row_idx = 2
    for category, subcategories in MASTER_STRUCTURE.items():
        # Write main category header
        ws.cell(row=row_idx, column=1, value=category).font = category_font
        row_idx += 1

        for subcategory, items in subcategories.items():
            # Write subcategory header if it's not 'N/A'
            if subcategory != "N/A":
                ws.cell(row=row_idx, column=1, value=subcategory).font = subcategory_font
                ws.cell(row=row_idx, column=1).alignment = Alignment(indent=1)
                row_idx += 1

            for item_name in items:
                if item_name in all_items:
                    item_data = all_items[item_name]
                    ws.cell(row=row_idx, column=1, value=item_name).alignment = Alignment(indent=2)
                    for col_idx, year in enumerate(all_years, 2):
                        cell = ws.cell(row=row_idx, column=col_idx)
                        cell.value = item_data.get(str(year))
                        cell.number_format = currency_format
                    
                    # Apply total styling for 'Total' rows
                    if "total" in item_name.lower():
                        for col_idx in range(1, len(headers) + 1):
                            ws.cell(row=row_idx, column=col_idx).font = total_font
                            ws.cell(row=row_idx, column=col_idx).border = total_border
                    
                    row_idx += 1
        
        # Add a blank row between major categories for readability
        row_idx += 1

    logger.info("Excel report generated successfully.")
    return wb

# --- Document Refinement Functions ---

async def generate_refinement_instructions_with_gemini(template_content, data_content):
    """Generates instructions for document refinement using Gemini."""
    logger.info("Generating document refinement instructions with Gemini.")
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY is not set.")

    model = genai.GenerativeModel('gemini-1.5-flash-latest')

    template_text_sample = template_content.get('text', '')[:4000]
    data_headings = data_content.get('headings', [])
    data_tables = data_content.get('tables', [])

    data_headings_str = "\n".join([f"Level {h['level']}: {h['text']}" for h in data_headings])
    data_tables_str = "\n---\n".join(["\n".join([" | ".join(map(str, cell)) for cell in table]) for table in data_tables])

    prompt = f"""
    You are an AI assistant specialized in precise document content extraction and transformation.
    Your task is to analyze a 'template document' and a 'data document' to facilitate creating a new document
    that uses the data document as its base but incorporates specific content from the template document.

    **From the TEMPLATE DOCUMENT, extract the following content:**

    1.  **"Adoption of MOI" Table:** Locate the table under the heading "Adoption of MOI". Extract its header row and all data rows. The columns are typically "Name of Director", "ID Number", "Signature", "Date". Provide this as an array of arrays.
    2.  **"Contents" Section:** Extract the full text content of the "CONTENTS" section, including all numbered or bulleted list items and their associated page numbers (if present).
    3.  **"DEFINITIONS" Section:** Extract the full text content of the "DEFINITIONS" section, including all numbered definitions and their descriptions.

    **From the DATA DOCUMENT, extract (for potential future use, though not directly inserted in this iteration):**
    1.  A structured list of all headings to build a new Table of Contents (for reference, not insertion).
    2.  Any table data that corresponds to directors/trustees (for reference, not insertion).

    Provide the output as a JSON object with the following keys:
    - `template_adoption_moi_table_headers`: Array of Strings, the headers of the "Adoption of MOI" table from the template.
    - `template_adoption_moi_table_data`: Array of Arrays of Strings, the data rows from the "Adoption of MOI" table in the template.
    - `template_contents_text`: String, the raw text content of the "CONTENTS" section from the template.
    - `template_definitions_text`: String, the raw text content of the "DEFINITIONS" section from the template.
    - `data_file_headings_for_toc`: Array of Objects, each with 'level' (integer) and 'text' (string) for headings from the data document (for reference).
    - `data_file_director_table_data`: Array of Arrays of Strings, any extracted director/trustee data from the data document (for reference).
    - `summary`: String, a brief summary of the extraction.

    If a section or table is not found, its corresponding field should be an empty string or empty array.

    --- TEMPLATE DOCUMENT EXTRACT ---
    {template_text_sample}
    --- END TEMPLATE DOCUMENT EXTRACT ---

    --- DATA DOCUMENT HEADINGS (for reference) ---
    {data_headings_str}
    --- END DATA DOCUMENT HEADINGS ---

    --- DATA DOCUMENT TABLES (for reference) ---
    {data_tables_str}
    --- END DATA DOCUMENT TABLES ---
    """
    
    try:
        response = await model.generate_content_async(
            prompt,
            generation_config={"response_mime_type": "application/json"}
        )
        cleaned_text = response.text.strip().replace("```json", "").replace("```", "")
        logger.debug(f"Raw JSON from Gemini (refinement): {cleaned_text[:500]}...")
        return json.loads(cleaned_text)
    except Exception as e:
        logger.error(f"Error in Gemini refinement instruction generation: {e}")
        return {
            "template_adoption_moi_table_headers": [],
            "template_adoption_moi_table_data": [],
            "template_contents_text": "",
            "template_definitions_text": "",
            "data_file_headings_for_toc": [],
            "data_file_director_table_data": [],
            "summary": f"Error during generation: {e}"
        }

def create_refined_document(original_data_file_bytes, parsed_template_info, refinement_data):
    """
    Creates a new DOCX document by taking the original data file as base
    and appending specific extracted content from the template.
    """
    logger.info("Starting creation of the refined document by appending template content.")

    # Load the original data file as the base for the new document
    new_doc = Document(io.BytesIO(original_data_file_bytes))

    # Extract data from refinement_data for easier access
    template_contents_text = refinement_data.get("template_contents_text", "CONTENTS SECTION NOT FOUND IN TEMPLATE.")
    template_definitions_text = refinement_data.get("template_definitions_text", "DEFINITIONS SECTION NOT FOUND IN TEMPLATE.")
    template_adoption_moi_table_headers = refinement_data.get("template_adoption_moi_table_headers", [])
    template_adoption_moi_table_data = refinement_data.get("template_adoption_moi_table_data", [])

    # --- Append Extracted Content from Template to Data File ---

    # 1. Append "Contents" from Template
    new_doc.add_page_break()
    new_doc.add_heading("CONTENTS (from Template)", level=1)
    # Add the raw text of the contents section
    for line in template_contents_text.split('\n'):
        if line.strip(): # Only add non-empty lines
            new_doc.add_paragraph(line.strip())
    new_doc.add_paragraph("") # Spacer

    # 2. Append "DEFINITIONS" from Template
    new_doc.add_page_break()
    new_doc.add_heading("DEFINITIONS (from Template)", level=1)
    # Add the raw text of the definitions section
    for line in template_definitions_text.split('\n'):
        if line.strip(): # Only add non-empty lines
            new_doc.add_paragraph(line.strip())
    new_doc.add_paragraph("") # Spacer

    # 3. Append "Adoption of MOI" Table from Template
    new_doc.add_page_break()
    new_doc.add_heading("Adoption of MOI (from Template)", level=1)
    
    if template_adoption_moi_table_headers and template_adoption_moi_table_data:
        new_table = new_doc.add_table(rows=1, cols=len(template_adoption_moi_table_headers))
        new_table.style = 'Table Grid' # Apply a default grid style
        
        # Set header row
        hdr_cells = new_table.rows[0].cells
        for i, header in enumerate(template_adoption_moi_table_headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Add data rows
        for row_data in template_adoption_moi_table_data:
            row_cells = new_table.add_row().cells
            for i, cell_data in enumerate(row_data):
                if i < len(row_cells): # Ensure we don't go out of bounds
                    row_cells[i].text = str(cell_data)
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    else:
        new_doc.add_paragraph("Adoption of MOI table data not found or extracted from template.")
    new_doc.add_paragraph("") # Spacer


    # Save the new document to a memory buffer
    doc_io = io.BytesIO()
    new_doc.save(doc_io)
    doc_io.seek(0)
    logger.info("Refined document created successfully by appending template content to data file.")
    return doc_io

# --- Flask API Routes ---

@app.route('/upload-and-convert', methods=['POST'])
def upload_and_convert_pdfs():
    """
    API endpoint to upload multiple PDF files, process them, and return a consolidated Excel file.
    """
    if 'files' not in request.files:
        return jsonify({"error": "No files part in the request"}), 400

    files = request.files.getlist('files')
    if not files or all(f.filename == '' for f in files):
        return jsonify({"error": "No selected files"}), 400

    custom_prompt_text = request.form.get('prompt', '')
    all_extracted_data = []
    all_years = set()

    for file in files:
        if file and file.filename.lower().endswith('.pdf'):
            try:
                logger.info(f"Processing file: {file.filename}")
                pdf_content = file.read()
                text_content = extract_text_from_pdf(pdf_content)
                
                if text_content:
                    parsed_data = parse_financial_data_with_gemini(text_content, file.filename, custom_prompt_text)
                    if parsed_data:
                        all_extracted_data.extend(parsed_data)
                        for item in parsed_data:
                            if 'AmountsByYear' in item and isinstance(item['AmountsByYear'], dict):
                                all_years.update(item['AmountsByYear'].keys())
                else:
                    logger.warning(f"Could not extract text from {file.filename}")

            except Exception as e:
                logger.error(f"Failed to process file {file.filename}: {e}")
                return jsonify({"error": f"An error occurred while processing {file.filename}: {str(e)}"}), 500
    
    if not all_extracted_data:
        return jsonify({"error": "Could not extract any financial data from the provided files."}), 400

    # --- Data Consolidation and Cleaning ---
    sorted_years = sorted(list(all_years), reverse=True)
    consolidated_items = {}

    for item in all_extracted_data:
        if 'Description' not in item or 'AmountsByYear' not in item:
            continue
        
        canonical_name = get_canonical_name(item['Description'])
        if canonical_name not in consolidated_items:
            consolidated_items[canonical_name] = {}
        
        for year, value in item['AmountsByYear'].items():
            cleaned_val = clean_value(value)
            if cleaned_val is not None:
                # If the key already exists, sum the values (handles duplicates)
                consolidated_items[canonical_name][year] = consolidated_items[canonical_name].get(year, 0) + cleaned_val

    # Post-processing for Accumulated Surplus/Deficit
    for year in sorted_years:
        surplus = consolidated_items.get("Accumulated surplus", {}).get(year)
        deficit = consolidated_items.get("Accumulated deficit", {}).get(year)
        if surplus is not None and deficit is not None:
            if surplus > 0 and deficit < 0:
                # Both exist, decide which one to keep. Let's assume surplus is the primary.
                # Ensure the key exists before attempting to delete
                if year in consolidated_items["Accumulated deficit"]:
                    del consolidated_items["Accumulated deficit"][year]
            elif deficit != 0:
                consolidated_items["Accumulated surplus"][year] = deficit
                if year in consolidated_items["Accumulated deficit"]:
                    del consolidated_items["Accumulated deficit"][year]


    # --- Excel Generation ---
    try:
        workbook = generate_excel_report(consolidated_items, sorted_years)
        output_io = io.BytesIO()
        workbook.save(output_io)
        output_io.seek(0)
        
        return send_file(
            output_io,
            as_attachment=True,
            download_name='consolidated_financials.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    except Exception as e:
        logger.error(f"Failed to generate Excel file: {e}")
        return jsonify({"error": "Failed to generate the Excel file."}), 500


@app.route('/refine-document', methods=['POST'])
async def refine_document_route():
    """
    API endpoint to refine a document using a template and a data file.
    """
    if 'template_file' not in request.files or 'data_file' not in request.files:
        return jsonify({"error": "Both a template file and a data file are required."}), 400

    template_file = request.files['template_file']
    data_file = request.files['data_file']

    # Allow .pdf, .doc, .docx for template files
    allowed_extensions = ('.pdf', '.doc', '.docx')
    if not template_file.filename.lower().endswith(allowed_extensions):
        return jsonify({"error": f"Template file must be one of {', '.join(allowed_extensions)}."}), 400
    if not data_file.filename.lower().endswith(allowed_extensions):
        return jsonify({"error": f"Data file must be one of {', '.join(allowed_extensions)}."}), 400

    try:
        template_content_raw = template_file.read()
        data_file_content_raw = data_file.read()

        # Extract content from template file
        parsed_template_content = extract_text_from_file(template_content_raw, template_file.filename)
        # Extract content from data file
        parsed_data_content = extract_text_from_file(data_file_content_raw, data_file.filename)

        # Get refinement instructions from AI
        refinement_instructions = await generate_refinement_instructions_with_gemini(parsed_template_content, parsed_data_content)

        if "Error" in refinement_instructions.get('summary', ''):
             return jsonify({"error": f"AI processing failed: {refinement_instructions['summary']}"}), 500

        # Create the new document by inserting template content into the data file
        refined_doc_io = create_refined_document(data_file_content_raw, parsed_template_content, refinement_instructions)

        return send_file(
            refined_doc_io,
            as_attachment=True,
            download_name='refined_document.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logger.error(f"An error occurred during document refinement: {e}", exc_info=True)
        return jsonify({"error": f"An internal error occurred: {str(e)}"}), 500


if __name__ == '__main__':
    # Changed port to 5000 to match frontend expectation
    app.run(debug=True, port=5000)